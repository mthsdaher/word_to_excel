import os

from docx import Document
from openpyxl import Workbook
from tqdm import tqdm

DATA_DIR = "./data"

# Create a new workbook and access the active sheet
wb = Workbook()
ws = wb.active

# Define column headers
headers = ["Invoice ID", "Total Products Purchased", "Subtotal", "Tax", "Total"]
ws.append(headers)


def parse_products_paragraph(paragraph):
    """Parses a raw product paragraph into a dictionary.

    Example:
        >>> product_paragram = "PRODUCTS\nCrampons:1\nParka:1\n"
        >>> parse_products_paragraph(product_paragram)
        {'Crampons': 1, 'Parka': 1}
    """
    products_dict = {}
    lines = paragraph.strip().split("\n")

    for line in lines[1:]:  # ignoring the PRODUCTS prefix, otherwise would be [0:]
        product_name, quantity = line.split(":") 
        products_dict[product_name] = int(quantity)
    return products_dict


def parse_footer_paragraph(paragraph):
    """Parses a raw product paragraph into a tuple.

    Example:
        >>> footer_paragraph = "SUBTOTAL:95.04\nTAX:12.36\nTOTAL:107.4"
        >>> parse_footer_paragraph(footer_paragraph)
        (95.04, 12.36, 107.4)
    """
    lines = paragraph.split("\n")
    subtotal = float(lines[0].split(":")[1])
    tax = float(lines[1].split(":")[1])
    total = float(lines[2].split(":")[1])
    return subtotal, tax, total


# Iterate through the docx files
files = list(sorted(os.listdir(DATA_DIR))) #sorted = ascending on alphabetic order - read the string from left to right
for file in tqdm(files, desc="Generating spreadsheet"): #tqdm is very usefull on this task. 
    if file.endswith(".docx"):
        filename = f"{DATA_DIR}/{file}"
        doc = Document(filename)

        # Extract information from the document
        invoice_id = doc.paragraphs[0].text

        products = parse_products_paragraph(doc.paragraphs[1].text)
        total_products = sum(products.values()) #the value from the dictionary --- Key -> Value

        subtotal, tax, total = parse_footer_paragraph(doc.paragraphs[2].text)

        # Add the information to the spreadsheet
        ws.append([invoice_id, total_products, subtotal, tax, total])

wb.save("output.xlsx")
wb.close()
