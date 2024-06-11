import os
import random

from docx import Document
from tqdm import tqdm

DATA_DIR = "./data"


def make_invoices(num_files):
    products = [
        "Parka",
        "Boots",
        "Snowshoes",
        "Climbing Rope",
        "Oxygen Tank",
        "Ice Pick",
        "Crampons",
    ]

    # Invoice loop
    for i in tqdm(range(num_files), desc="Generating invoices"):
        # Create Randomized invoice
        invoice_num = "100" + str(i).zfill(4)
        product_dict = {}

        for _ in range(random.randint(1, 10)):
            product = products[random.randint(0, len(products) - 1)]
            if product in product_dict:
                product_dict[product] += 1
            else:
                product_dict[product] = 1

        subtot = round(random.random() * 10 ** (random.randint(3, 4)), 2)
        tax = round(subtot * 0.13, 2)
        total = round(subtot + tax, 2)

        # Create doc from random invoice
        doc = Document()
        doc.add_heading("INV" + invoice_num)
        prod = doc.add_paragraph("PRODUCTS\n")
        for key in product_dict.keys():
            prod.add_run(f"{key}:{product_dict[key]}\n")
        doc.add_paragraph(f"SUBTOTAL:{subtot}\nTAX:{tax}\nTOTAL:{total}")
        doc.save(f"{DATA_DIR}/INV{invoice_num}.docx")


if not os.path.exists(DATA_DIR):
    os.makedirs(DATA_DIR)

make_invoices(200)
